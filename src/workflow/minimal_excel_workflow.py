from __future__ import annotations

import asyncio
import html
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

# Allow direct run: python src/workflow/minimal_excel_workflow.py
SRC_ROOT = Path(__file__).resolve().parents[1]
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from deli_api import CaseHit, DeliLegalClient, LawHit, init_case_client
from dotenv import load_dotenv
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from pydantic import BaseModel, ConfigDict, Field


# =========================
# Debug Config (edit here)
# =========================
QUERY = "上班途中车祸工伤案例"
TARGET_CASE_COUNT = 50
REWRITE_QUERY_COUNT = 3
PAGE_SIZE = 5  # Deli API 每页最多 5 条
MAX_CONCURRENCY = 5
DOTENV_PATH = Path(".env")
WORD_SUBMITTER = "xxxxxx"


HEADERS = [
    "审理法院",
    "案号",
    "案由",
    "法院审理程序",
    "法院层级",
    "法院认为（精简后）",
    "裁判结果",
    "主要原因",
    "关键裁判结论-权利类",
    "关键裁判结论-金额类",
    "关键裁判结论-行为类",
    "法律依据（法律名称+条文）",
]


class KeyConclusions(BaseModel):
    权利类: str = Field(..., min_length=1, description="权利义务方向的结论")
    金额类: str = Field(..., min_length=1, description="金额方向的结论")
    行为类: str = Field(..., min_length=1, description="行为义务方向的结论")


class CaseExtractResult(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    案由: str = Field(..., min_length=1, description="案件案由")
    法院审理程序: str = Field(..., min_length=1, description="例如：一审/二审/再审")
    法院层级: str = Field(..., min_length=1, description="例如：基层法院/中院/高院/最高法院")
    法院认为_精简后: str = Field(
        ...,
        alias="法院认为（精简后）",
        min_length=1,
        description="提炼后的法院说理，1-4条",
    )
    裁判结果: str = Field(..., min_length=1, description="判决主文/裁判结果")
    主要原因: str = Field(..., min_length=1, description="主要裁判原因，1-3句")
    关键裁判结论: KeyConclusions = Field(..., description="按三类输出的关键裁判结论")


class QueryRewriteResult(BaseModel):
    改写查询: list[str] = Field(
        ...,
        min_length=REWRITE_QUERY_COUNT,
        max_length=REWRITE_QUERY_COUNT,
        description=f"必须提供{REWRITE_QUERY_COUNT}条不同的法律检索语句",
    )


class LawQueryRewriteResult(BaseModel):
    法律检索查询: list[str] = Field(
        ...,
        min_length=REWRITE_QUERY_COUNT,
        max_length=REWRITE_QUERY_COUNT,
        description="用于检索法律条文的专业化短查询语句",
    )


class MinimalCaseRow(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    审理法院: str = Field(..., min_length=1)
    案号: str = Field(..., min_length=1)
    案由: str = Field(..., min_length=1)
    法院审理程序: str = Field(..., min_length=1)
    法院层级: str = Field(..., min_length=1)
    法院认为_精简后: str = Field(..., alias="法院认为（精简后）", min_length=1)
    裁判结果: str = Field(..., min_length=1)
    主要原因: str = Field(..., min_length=1)
    关键裁判结论_权利类: str = Field(..., min_length=1)
    关键裁判结论_金额类: str = Field(..., min_length=1)
    关键裁判结论_行为类: str = Field(..., min_length=1)
    法律依据: str = Field(..., alias="法律依据（法律名称+条文）", min_length=1)

    def to_excel_row(self) -> list[str]:
        return [
            self.审理法院,
            self.案号,
            self.案由,
            self.法院审理程序,
            self.法院层级,
            self.法院认为_精简后,
            self.裁判结果,
            self.主要原因,
            self.关键裁判结论_权利类,
            self.关键裁判结论_金额类,
            self.关键裁判结论_行为类,
            self.法律依据,
        ]


class OpeningContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    报告标题: str = Field(..., min_length=1, description="例如：案涉争议问题相关检索报告")
    开篇亮明观点: str = Field(
        ...,
        min_length=1,
        description="必须同时写明核心争议问题与本方明确法律结论，旗帜鲜明",
    )
    争议焦点段: str = Field(..., min_length=1, description="围绕本案争议焦点进行凝练表述")
    类案检索结论段: str = Field(..., min_length=1, description="提炼检索结论并服务本案主张")


class SectionOneContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    第一部分标题: str = Field(..., min_length=1, description="固定格式示例：一、【待决案件案情简述】")
    第一部分正文: str = Field(
        ...,
        min_length=1,
        description="围绕待决案件事实与为何需要类案比较进行正式法律文书写作",
    )


class SectionTwoContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    第二部分标题: str = Field(..., min_length=1, description="固定格式示例：二、【类案基本事实概括】")
    检索方法段: str = Field(
        ...,
        min_length=1,
        description="说明检索平台、关键词、检索法院、检索时间等方法信息",
    )
    类案检索情况段: str = Field(
        ...,
        min_length=1,
        description="总结检索规模、裁判结果分布与关键统计结论",
    )
    本案关联性段: str = Field(
        ...,
        min_length=1,
        description="提炼类案与本案直接相关的事实或法律争点关联",
    )


class SectionThreeTableRow(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    中级人民法院案例: str = Field(..., min_length=1, description="法院名称+案号")
    案由: str = Field(..., min_length=1, description="该类案案由")
    裁判要点与理由: str = Field(..., min_length=1, description="凝练裁判要点及理由")


class SectionThreeViewpoint(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    观点标题: str = Field(..., min_length=1, description="例如：观点一")
    观点总结: str = Field(..., min_length=1, description="围绕该观点的明确结论")
    表格标题: str = Field(..., min_length=1, description="例如：表格1")
    裁判要点表格: list[SectionThreeTableRow] = Field(..., min_length=1, description="该观点下的类案表格")


class SectionThreeContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    第三部分标题: str = Field(..., min_length=1, description="固定格式示例：三、【类案检索裁判要点】")
    第三部分引言段: str = Field(..., min_length=1, description="说明第三部分摘录口径与参考价值")
    分类说明段: str = Field(..., min_length=1, description="说明按何种观点进行分类")
    观点列表: list[SectionThreeViewpoint] = Field(..., min_length=1, description="至少一个观点分组")


class SectionFourLawItem(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    法规名称与条款: str = Field(..., min_length=1, description="例如：《建设工程司法解释（一）》第二十六条")
    条文原文: str = Field(..., min_length=1, description="对应条款原文，保留法律文书风格")
    条文要点列表: list[str] = Field(..., min_length=1, description="按（一）（二）等形式展开的要点")


class SectionFourContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    第四部分标题: str = Field(..., min_length=1, description="固定格式示例：四、【相关法律法规原文：法律、司法解释】")
    第四部分引言段: str = Field(..., min_length=1, description="说明本部分收录条文与适用目的")
    法规条文列表: list[SectionFourLawItem] = Field(..., min_length=1, description="与本案争议直接相关的法律条文")


class SectionFiveContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    第五部分标题: str = Field(..., min_length=1, description="固定格式示例：五、【类案检索结果分析】")
    分析结论列表: list[str] = Field(..., min_length=3, max_length=3, description="按①②③输出三条结论性分析")


class AttachmentCaseSelection(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    选中序号: list[int] = Field(..., min_length=1, description="从候选案例中选出最适合作为附件的序号，优先6-7个")


class WordReportContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    报告标题: str = Field(..., min_length=1, description="例如：案涉争议问题相关检索报告")
    开篇亮明观点: str = Field(
        ...,
        min_length=1,
        description="必须同时包含核心争议问题与本方明确法律结论，并出现‘检索结果显示，生效裁判均一致认为’",
    )
    待决案件案情简述: str = Field(..., min_length=1, description="聚焦待决案件与争议焦点相关事实的简述")
    第一部分_类案基本事实概括: str = Field(..., min_length=1, description="仅概括与本案争议焦点直接相关且相似的类案事实")
    第一部分_检索方法: str = Field(..., min_length=1, description="写明检索平台、关键词、检索法院、检索时间")
    第一部分_检索情况: str = Field(..., min_length=1, description="概括类案检索数量与结果分布")
    第一部分_本案关联性: str = Field(..., min_length=1, description="明确类案与本案直接关联")
    第二部分_类案核心裁判要旨: str = Field(..., min_length=1, description="仅围绕本方核心主张提炼裁判要旨")
    第二部分_观点总结列表: list[str] = Field(..., min_length=1, description="按观点分组总结关键裁判要旨")
    第三部分_关联性与最终结论: str = Field(..., min_length=1, description="解读类案与本案关联性并给出最终结论")
    第四部分_相关法律法规原文: list[str] = Field(..., min_length=1, description="与争议问题直接相关的法律法规和司法解释原文")
    第五部分_结果分析: list[str] = Field(..., min_length=3, max_length=3, description="三条结论性分析，分别对应相似性、参照观点、效力层级")
    应当参照类案: list[str] = Field(..., min_length=1, description="应当参照类案及依据")
    可以参考类案: list[str] = Field(..., min_length=1, description="可以参考类案及依据")


def load_env() -> None:
    load_dotenv(dotenv_path=DOTENV_PATH, override=False)


def init_llm() -> ChatOpenAI:
    api_key = os.getenv("LLM_API_KEY")
    model = os.getenv("LLM_MODEL")
    base_url = os.getenv("LLM_BASE_URL")
    if not api_key or not model or not base_url:
        missing = [
            key
            for key, value in (
                ("LLM_API_KEY", api_key),
                ("LLM_MODEL", model),
                ("LLM_BASE_URL", base_url),
            )
            if not value
        ]
        raise RuntimeError(f"Missing required LLM env vars: {', '.join(missing)}")

    return ChatOpenAI(
        model=model,
        api_key=api_key,
        temperature=0.1,
        base_url=base_url,
    )


def build_case_dedup_key(hit: CaseHit) -> str:
    case_no = (hit.case_no or "").replace("（", "(").replace("）", ")").replace(" ", "")
    if not case_no:
        raise ValueError(f"Case missing case_no: {hit.title}")
    return case_no


def build_law_basis(hits: list[LawHit]) -> str:
    parts: list[str] = []
    for hit in hits:
        law_name = re.sub(r"<[^>]+>", "", html.unescape(hit.law_name or "")).strip()
        article_no = re.sub(r"<[^>]+>", "", html.unescape(hit.article_no or "")).strip()
        if law_name and article_no:
            parts.append(f"{law_name}{article_no}")
        if len(parts) >= 2:
            break
    return "；".join(parts)


async def rewrite_queries(rewriter: Any, query: str) -> list[str]:
    if not query.strip():
        raise ValueError("QUERY must not be empty.")
    schema_text = json.dumps(QueryRewriteResult.model_json_schema(), ensure_ascii=False, indent=2)
    result: QueryRewriteResult = await rewriter.ainvoke(
        [
            SystemMessage(content="你是法律检索改写助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"用户原始查询：{query}\n"
                    f"请输出{REWRITE_QUERY_COUNT}条不同的专业法律检索语句。\n"
                    "要求：不加解释、不加编号、不使用原句。\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )
    rewritten = [item.strip() for item in result.改写查询]
    if len(set(rewritten)) != len(rewritten):
        raise ValueError("Rewritten queries contain duplicates.")
    if any(item == query for item in rewritten):
        raise ValueError("Rewritten queries contain the original query.")
    return rewritten


async def fetch_case_hits(case_client: Any, query: str, target_count: int) -> list[CaseHit]:
    if target_count <= 0:
        raise ValueError("TARGET_CASE_COUNT must be greater than 0.")
    if PAGE_SIZE < 1 or PAGE_SIZE > 5:
        raise ValueError("PAGE_SIZE must be between 1 and 5.")
    page_no = 1
    hits: list[CaseHit] = []
    seen: set[str] = set()
    while len(hits) < target_count:
        page_hits = await asyncio.to_thread(
            case_client.search_cases,
            query,
            page_no=page_no,
            page_size=PAGE_SIZE,
        )
        if not page_hits:
            break
        for hit in page_hits:
            key = build_case_dedup_key(hit)
            if key in seen:
                continue
            seen.add(key)
            hits.append(hit)
            if len(hits) >= target_count:
                break
        if len(page_hits) < PAGE_SIZE:
            break
        page_no += 1
    return hits


async def extract_case_fields(extractor: Any, hit: CaseHit) -> CaseExtractResult:
    if not hit.title.strip():
        raise ValueError("Case title is empty.")
    if not hit.content.strip():
        raise ValueError(f"Case content is empty: {hit.title}")
    schema_text = json.dumps(CaseExtractResult.model_json_schema(), ensure_ascii=False, indent=2)
    return await extractor.ainvoke(
        [
            SystemMessage(content="你是法律要素抽取助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"案件标题：{hit.title}\n"
                    f"审理法院：{hit.court_name}\n"
                    f"案号：{hit.case_no}\n"
                    f"案由：{hit.cause}\n"
                    f"审理程序（元数据）：{hit.level_of_trial}\n"
                    f"案件类型（元数据）：{hit.case_type}\n\n"
                    f"裁判文书原文：\n{hit.content}\n\n"
                    "请基于以上内容提取结果。\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def rewrite_law_queries(
    law_rewriter: Any,
    hit: CaseHit,
    extracted: CaseExtractResult,
) -> list[str]:
    schema_text = json.dumps(LawQueryRewriteResult.model_json_schema(), ensure_ascii=False, indent=2)
    result: LawQueryRewriteResult = await law_rewriter.ainvoke(
        [
            SystemMessage(content="你是法律条文检索助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    "请将以下案件信息改写为可用于法条检索的专业查询语句。\n"
                    "要求：每条查询是短语，不要解释，不要编号，不要口语化长句。\n"
                    f"案件标题：{hit.title}\n"
                    f"案由：{extracted.案由}\n"
                    f"主要原因：{extracted.主要原因}\n"
                    f"裁判结果：{extracted.裁判结果}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )

    rewritten = [item.strip() for item in result.法律检索查询 if item.strip()]
    base_queries = [
        f"{hit.title} {extracted.主要原因}".strip(),
        extracted.案由.strip(),
        (hit.cause or "").strip(),
        extracted.主要原因.strip(),
    ]
    merged: list[str] = []
    seen: set[str] = set()
    for q in [*rewritten, *base_queries]:
        if q and q not in seen:
            seen.add(q)
            merged.append(q)
    return merged


async def process_hit(
    extractor: Any,
    law_rewriter: Any,
    law_client: DeliLegalClient,
    hit: CaseHit,
    semaphore: asyncio.Semaphore,
) -> MinimalCaseRow:
    async with semaphore:
        if not hit.court_name or not hit.case_no:
            raise ValueError(f"Case metadata missing court_name/case_no: {hit.title}")

        extracted = await extract_case_fields(extractor, hit)
        law_queries = await rewrite_law_queries(law_rewriter, hit, extracted)
        law_basis = ""
        for query in law_queries:
            q = query.strip()
            if not q:
                continue
            law_hits = await asyncio.to_thread(
                law_client.search_laws,
                q,
                page_size=3,
                field_name="semantic",
            )
            law_basis = build_law_basis(law_hits)
            if law_basis:
                break
        if not law_basis:
            law_basis = "未检索到明确法律条文"

        return MinimalCaseRow(
            审理法院=hit.court_name.strip(),
            案号=hit.case_no.strip(),
            案由=extracted.案由.strip(),
            法院审理程序=extracted.法院审理程序.strip(),
            法院层级=extracted.法院层级.strip(),
            法院认为_精简后=extracted.法院认为_精简后.strip(),
            裁判结果=extracted.裁判结果.strip(),
            主要原因=extracted.主要原因.strip(),
            关键裁判结论_权利类=extracted.关键裁判结论.权利类.strip(),
            关键裁判结论_金额类=extracted.关键裁判结论.金额类.strip(),
            关键裁判结论_行为类=extracted.关键裁判结论.行为类.strip(),
            法律依据=law_basis.strip(),
        )


async def build_rows() -> tuple[list[MinimalCaseRow], list[CaseHit]]:
    with init_case_client() as case_client, DeliLegalClient.from_env(DOTENV_PATH) as law_client:
        llm = init_llm()
        rewriter = llm.with_structured_output(QueryRewriteResult, method="function_calling")
        law_rewriter = llm.with_structured_output(LawQueryRewriteResult, method="function_calling")
        extractor = llm.with_structured_output(CaseExtractResult, method="function_calling")

        rewritten_queries = await rewrite_queries(rewriter, QUERY)
        all_queries = [QUERY, *rewritten_queries]

        merged_hits: list[CaseHit] = []
        seen: set[str] = set()
        for query in all_queries:
            hits = await fetch_case_hits(case_client, query, TARGET_CASE_COUNT)
            for hit in hits:
                key = build_case_dedup_key(hit)
                if key in seen:
                    continue
                seen.add(key)
                merged_hits.append(hit)
                if len(merged_hits) >= TARGET_CASE_COUNT:
                    break
            if len(merged_hits) >= TARGET_CASE_COUNT:
                break

        semaphore = asyncio.Semaphore(MAX_CONCURRENCY)
        tasks = [process_hit(extractor, law_rewriter, law_client, hit, semaphore) for hit in merged_hits]
        rows = await asyncio.gather(*tasks)
        return rows, merged_hits


async def generate_opening_content(query: str, rows: list[MinimalCaseRow]) -> OpeningContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    opening_writer = llm.with_structured_output(OpeningContent, method="function_calling")
    schema_text = json.dumps(OpeningContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows[:12]], ensure_ascii=False, indent=2)

    return await opening_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    "请撰写报告开篇，要求严格遵守：\n"
                    "1) 开篇必须旗帜鲜明亮明观点：待论证核心争议问题 + 本方明确法律结论。\n"
                    "2) 不得只写“提交报告供参考”，必须直接回应“要解决什么问题、实现什么法律效果”。\n"
                    "3) 文风正式、法庭提交场景可直接使用。\n"
                    "4) 开篇亮明观点段中必须出现“检索结果显示，生效裁判均一致认为”。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_section_one_content(
    query: str,
    opening: OpeningContent,
    rows: list[MinimalCaseRow],
) -> SectionOneContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    section_writer = llm.with_structured_output(SectionOneContent, method="function_calling")
    schema_text = json.dumps(SectionOneContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows[:12]], ensure_ascii=False, indent=2)

    return await section_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    f"开篇观点：{opening.开篇亮明观点}\n"
                    "请撰写报告第一部分，要求：\n"
                    "1) 第一部分标题必须为“一、【待决案件案情简述】”。\n"
                    "2) 正文须说明待决案件核心事实背景，并明确为何需要将待决案件与类案进行比较。\n"
                    "3) 文风庄重、书面化，适合直接提交合议庭。\n"
                    "4) 不要使用占位符，不要写空泛套话。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_section_two_content(
    query: str,
    section_one: SectionOneContent,
    rows: list[MinimalCaseRow],
) -> SectionTwoContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    section_writer = llm.with_structured_output(SectionTwoContent, method="function_calling")
    schema_text = json.dumps(SectionTwoContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)

    return await section_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    f"第一部分内容：{section_one.第一部分正文}\n"
                    "请撰写报告第二部分，要求：\n"
                    "1) 第二部分标题必须为“二、【类案基本事实概括】”。\n"
                    "2) 检索方法段必须覆盖：检索平台、检索关键词、检索法院、检索时间。\n"
                    "3) 类案检索情况段必须概括检索数量与结果分布，语气客观正式。\n"
                    "4) 本案关联性段必须明确指出类案与本案直接相关的法律争点。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_section_three_content(
    query: str,
    section_two: SectionTwoContent,
    rows: list[MinimalCaseRow],
) -> SectionThreeContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    section_writer = llm.with_structured_output(SectionThreeContent, method="function_calling")
    schema_text = json.dumps(SectionThreeContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)

    return await section_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    f"第二部分摘要：{section_two.类案检索情况段}\n"
                    "请撰写报告第三部分，要求：\n"
                    "1) 第三部分标题必须为“三、【类案检索裁判要点】”。\n"
                    "2) 引言段应说明：参考规范、摘录字段、并以表格形式展示。\n"
                    "3) 必须按观点分类输出，每个观点都要有“观点总结”和“裁判要点表格”。\n"
                    "4) 表格字段固定为：中级人民法院案例、案由、裁判要点与理由。\n"
                    "5) 每个观点至少给出1行表格数据，语言正式，适合提交合议庭。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_section_four_content(
    query: str,
    section_three: SectionThreeContent,
    rows: list[MinimalCaseRow],
) -> SectionFourContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    section_writer = llm.with_structured_output(SectionFourContent, method="function_calling")
    schema_text = json.dumps(SectionFourContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)

    return await section_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    f"第三部分摘要：{section_three.第三部分引言段}\n"
                    "请撰写报告第四部分，要求：\n"
                    "1) 第四部分标题必须为“四、【相关法律法规原文：法律、司法解释】”。\n"
                    "2) 每条法规需包含“法规名称与条款”“条文原文”“条文要点列表”。\n"
                    "3) 条文要点列表按规范法律文书写法输出，内容与本案争议直接相关。\n"
                    "4) 文风正式，适合法庭提交。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_section_five_content(
    query: str,
    section_three: SectionThreeContent,
    section_four: SectionFourContent,
    rows: list[MinimalCaseRow],
) -> SectionFiveContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    section_writer = llm.with_structured_output(SectionFiveContent, method="function_calling")
    schema_text = json.dumps(SectionFiveContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)

    return await section_writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    f"第三部分摘要：{section_three.第三部分引言段}\n"
                    f"第四部分摘要：{section_four.第四部分引言段}\n"
                    "请撰写报告第五部分，要求：\n"
                    "1) 第五部分标题必须为“五、【类案检索结果分析】”。\n"
                    "2) 分析结论列表必须恰好3条，分别对应①②③。\n"
                    "3) 第1条强调类案与本案核心事实相似点，强化类案可适用性。\n"
                    "4) 第2条明确提交合议庭参考的核心裁判观点及其对应事实认定理由。\n"
                    "5) 第3条必须从“结论预断（关于统一法律适用和强类案检索的指导意义）”角度总结，明确哪些属于法院应当参照的类案、哪些属于可参考类案。\n"
                    "6) 文风正式，适合法庭提交。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def generate_word_report_content(query: str, rows: list[MinimalCaseRow]) -> WordReportContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    llm = init_llm()
    writer = llm.with_structured_output(WordReportContent, method="function_calling")
    schema_text = json.dumps(WordReportContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)

    return await writer.ainvoke(
        [
            SystemMessage(content="你是法律文书写作助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    "请根据以下撰写要点输出类案检索报告正文内容：\n"
                    "【要点一】开篇必须旗帜鲜明亮明观点：待论证核心争议问题 + 本方明确法律结论。\n"
                    "【要点二】采用固定三段式：第一部分概括类案基本事实；第二部分总结类案核心裁判要旨；第三部分解读类案与本案关联性并得出最终结论。\n"
                    "【要点三】类案事实梳理详略得当并高度绑定本案：仅写与争议焦点直接相关事实，明确核心相似点。\n"
                    "【要点四】裁判要旨必须高度聚焦本方核心主张：无关观点剔除，仅可少量摘抄关键判词。\n"
                    "【要点五】结论清晰有力：重申相似点；明确提请合议庭参考/参照的核心观点及理由；结合《关于统一法律适用加强类案检索的指导意见》区分应当参照类案与可以参考类案。\n"
                    "并且必须覆盖模板中的具体板块：待决案件案情简述、类案基本事实概括（含检索方法/检索情况/本案关联性）、类案核心裁判要旨（含观点总结）、相关法律法规原文、类案检索结果分析。\n"
                    "额外硬性要求：\n"
                    "1) 开篇亮明观点中必须出现“检索结果显示，生效裁判均一致认为”。\n"
                    "2) 三段正文必须分别对应第一部分/第二部分/第三部分。\n"
                    "3) 文风正式，可直接提交合议庭。\n"
                    "4) 严禁输出空段落或仅含空白字符的段落。\n"
                    "5) 第五部分结果分析必须恰好三条，分别对应：核心事实相似点、提请参考/参照的核心观点与理由、应当参照/可以参考效力层级区分。\n"
                    f"类案样本（用于归纳）：\n{rows_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )


async def select_attachment_cases(
    query: str,
    rows: list[MinimalCaseRow],
    case_hits: list[CaseHit],
) -> list[CaseHit]:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not case_hits:
        raise ValueError("case_hits must not be empty.")
    if len(rows) != len(case_hits):
        raise ValueError("rows and case_hits length mismatch.")

    llm = init_llm()
    selector = llm.with_structured_output(AttachmentCaseSelection, method="function_calling")
    schema_text = json.dumps(AttachmentCaseSelection.model_json_schema(), ensure_ascii=False, indent=2)

    candidates: list[dict[str, Any]] = []
    for idx, (row, hit) in enumerate(zip(rows, case_hits), start=1):
        candidates.append(
            {
                "序号": idx,
                "审理法院": hit.court_name,
                "案号": hit.case_no,
                "标题": hit.title,
                "案由": row.案由,
                "裁判结果": row.裁判结果,
                "主要原因": row.主要原因,
            }
        )
    candidates_text = json.dumps(candidates, ensure_ascii=False, indent=2)

    result: AttachmentCaseSelection = await selector.ainvoke(
        [
            SystemMessage(content="你是法律案例筛选助手，输出必须严格符合给定结构。"),
            HumanMessage(
                content=(
                    f"本案检索主题：{query}\n"
                    "请从候选案例中筛选6-7个用于报告第六部分附件，筛选标准：\n"
                    "1) 优先选择高级人民法院或最高人民法院案例；\n"
                    "2) 其次可选择中级人民法院案例，但必须与本案核心事实和争点高度相似；\n"
                    "3) 不选基层法院且相关性弱的案例；\n"
                    "4) 数量上优先选择6-7个，但并非硬性要求，应以质量优先；\n"
                    "5) 输出仅包含“选中序号”数组。\n"
                    f"候选案例如下：\n{candidates_text}\n"
                    f"结构定义如下：\n{schema_text}"
                )
            ),
        ]
    )

    selected_hits: list[CaseHit] = []
    seen: set[int] = set()
    for i in result.选中序号:
        if i < 1 or i > len(case_hits):
            raise ValueError(f"Attachment case index out of range: {i}")
        if i in seen:
            continue
        seen.add(i)
        selected_hits.append(case_hits[i - 1])
    if not selected_hits:
        raise ValueError("Selected attachment case count must be greater than 0.")
    return selected_hits


def save_excel_sync(rows: list[MinimalCaseRow]) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "类案检索结果"

    ws.append(HEADERS)
    for row in rows:
        ws.append(row.to_excel_row())

    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    widths = [16, 24, 20, 16, 14, 40, 34, 28, 22, 22, 22, 40]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=12):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    output_path = Path("outputs") / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    return output_path


async def save_excel(rows: list[MinimalCaseRow]) -> Path:
    return await asyncio.to_thread(save_excel_sync, rows)


async def save_word(
    case_title: str,
    attachment_case_hits: list[CaseHit],
    report_content: WordReportContent,
    report_title: str = "案涉争议问题检索报告",
    submitter: str = WORD_SUBMITTER,
    report_date: str | None = None,
) -> Path:
    if not case_title:
        raise ValueError("case_title must not be empty.")

    def _build_and_save() -> Path:
        date_text = report_date or datetime.now().strftime("%Y年%m月")
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

        style = doc.styles["Normal"]
        style.font.name = "仿宋"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        style.font.size = Pt(14)

        label_color = RGBColor(47, 84, 150)

        def add_spacer(count: int) -> None:
            for _ in range(count):
                doc.add_paragraph("")

        def add_center_line(label: str, value: str, size: int) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r1 = p.add_run(f"【{label}】")
            r1.font.name = "仿宋"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r1.font.bold = True
            r1.font.size = Pt(size)
            r1.font.color.rgb = label_color

            r2 = p.add_run(value)
            r2.font.name = "仿宋"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r2.font.bold = True
            r2.font.size = Pt(size)

        def add_body_paragraph(text: str) -> None:
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if not lines:
                return
            for line in lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(28)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                run.font.size = Pt(12)

        def add_body_field_label(label: str) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(label)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(14)
            run.font.bold = True

        def add_body_field_value(label: str, value: str) -> None:
            value_lines = [line.strip() for line in value.splitlines() if line.strip()]
            if not value_lines:
                return
            for idx, line in enumerate(value_lines):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(28)
                p.paragraph_format.line_spacing = 1.5
                if idx == 0:
                    r1 = p.add_run(label)
                    r1.font.name = "仿宋"
                    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                    r1.font.size = Pt(12)
                    r1.font.bold = True
                r2 = p.add_run(line if idx > 0 else line)
                r2.font.name = "仿宋"
                r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                r2.font.size = Pt(12)

        def add_section_heading(text: str) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = label_color

        def clean_attachment_text(text: str) -> str:
            cleaned_text = html.unescape(text)
            cleaned_text = cleaned_text.replace("\r\n", "\n").replace("\r", "\n")
            cleaned_text = re.sub(r"<[^>]+>", "", cleaned_text)
            cleaned_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned_text)
            cleaned_text = re.sub(r"__(.*?)__", r"\1", cleaned_text)
            cleaned_text = re.sub(r"`([^`]+)`", r"\1", cleaned_text)
            cleaned_text = re.sub(r"\\</?em>", "", cleaned_text)
            cleaned_text = re.sub(r"[ \t\f\v]+", " ", cleaned_text)
            cleaned_text = re.sub(r"\n\s*\n+", "\n", cleaned_text)
            return cleaned_text.strip()

        def add_attachment_title(text: str) -> None:
            title = clean_attachment_text(text)
            if not title:
                return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(title)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.underline = True

        def add_attachment_paragraph(text: str) -> None:
            cleaned_text = clean_attachment_text(text)
            lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
            if not lines:
                return
            for line in lines:
                p = doc.add_paragraph()
                is_case_no_line = bool(re.match(r"^（?\d{4}）", line))
                is_signature_line = bool(
                    re.match(r"^(审判长|审判员|书记员|法官助理|代理审判员)", line)
                )
                is_date_footer_line = bool(
                    re.match(r"^(二〇|二○|二O|\d{4}年)", line)
                )
                if is_case_no_line or is_signature_line or is_date_footer_line:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.first_line_indent = Pt(0)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                run.font.size = Pt(10.5)

        def add_attachment_field(label: str, value: str, right_aligned: bool = False) -> None:
            label_text = clean_attachment_text(label)
            value_text = clean_attachment_text(value)
            if not label_text and not value_text:
                return
            p = doc.add_paragraph()
            if right_aligned:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.first_line_indent = Pt(0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.line_spacing = 1.5

            r1 = p.add_run(label_text)
            r1.font.name = "仿宋"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r1.font.size = Pt(10.5)
            r1.font.bold = True

            r2 = p.add_run(value_text)
            r2.font.name = "仿宋"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r2.font.size = Pt(10.5)

        add_spacer(1)
        add_center_line("案由", case_title, 16)
        add_spacer(4)
        add_center_line("标题", report_title, 22)
        add_spacer(7)
        add_center_line("署名", f"提交人：{submitter}", 14)
        add_spacer(1)
        add_center_line("日期", date_text, 14)

        doc.add_page_break()
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(report_content.报告标题)
        title_run.font.name = "仿宋"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        title_run.font.bold = True
        title_run.font.size = Pt(22)

        add_body_paragraph("尊敬的合议庭：")
        add_body_paragraph(report_content.开篇亮明观点)
        add_section_heading("【待决案件案情简述】")
        add_body_paragraph(report_content.待决案件案情简述)
        add_section_heading("一、【类案基本事实概括】")
        add_body_paragraph(report_content.第一部分_类案基本事实概括)
        add_body_field_value("【写明检索方法：包括检索平台、检索关键词、检索法院、检索时间】", report_content.第一部分_检索方法)
        add_body_field_value("【总结类案检索情况】", report_content.第一部分_检索情况)
        add_body_field_value("【要提及类案与本案直接的关联性】", report_content.第一部分_本案关联性)
        add_section_heading("二、【类案核心裁判要旨】")
        add_body_paragraph(report_content.第二部分_类案核心裁判要旨)
        for idx, view in enumerate(report_content.第二部分_观点总结列表, start=1):
            add_body_field_value(f"【观点{idx}】", view)
        add_section_heading("三、【类案与本案关联性及最终结论】")
        add_body_paragraph(report_content.第三部分_关联性与最终结论)
        add_section_heading("四、【相关法律法规原文：法律、司法解释】")
        for idx, law_text in enumerate(report_content.第四部分_相关法律法规原文, start=1):
            add_body_paragraph(f"{idx}、{law_text}")
        add_section_heading("五、【类案检索结果分析】")
        for idx, item in enumerate(report_content.第五部分_结果分析, start=1):
            add_body_paragraph(f"{idx}、{item}")
        add_body_field_label("应当参照类案：")
        for idx, item in enumerate(report_content.应当参照类案, start=1):
            add_body_paragraph(f"{idx}、{item}")
        add_body_field_label("可以参考类案：")
        for idx, item in enumerate(report_content.可以参考类案, start=1):
            add_body_paragraph(f"{idx}、{item}")

        doc.add_page_break()
        add_section_heading("六、【附件：类案检索原文】")
        add_attachment_paragraph("案例原文：")
        for idx, hit in enumerate(attachment_case_hits, start=1):
            if idx > 1:
                doc.add_page_break()
            add_attachment_title(f"{hit.title}")
            add_attachment_field("审理法院：", f" {hit.court_name or ''}")
            add_attachment_field("（案号）", f"{hit.case_no or ''}", right_aligned=True)
            add_attachment_field("案由：", f" {hit.cause or ''}")
            add_attachment_field("裁判日期：", f" {hit.case_date or ''}")
            add_attachment_paragraph(hit.content)

        output_path = Path("outputs") / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    return await asyncio.to_thread(
        _build_and_save,
    )


async def main() -> None:
    load_env()
    rows, case_hits = await build_rows()
    excel_path = await save_excel(rows)
    report_content = await generate_word_report_content(QUERY, rows)
    attachment_case_hits = await select_attachment_cases(QUERY, rows, case_hits)
    word_report_path = await save_word(
        case_title=QUERY,
        attachment_case_hits=attachment_case_hits,
        report_content=report_content,
    )
    print(f"Generated Excel: {excel_path.resolve()}")
    print(f"Generated Word Report: {word_report_path.resolve()}")
    print(f"Rows: {len(rows)}")
    print("Columns:", " | ".join(HEADERS))


if __name__ == "__main__":
    asyncio.run(main())
