from __future__ import annotations

import asyncio
import html
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from deli_api import CaseHit
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel, ConfigDict, Field

from case_table_adapter import CaseTableRow, build_case_table_rows_for_viewpoint
from excel_workflow import QUERY, MinimalCaseRow, build_rows, init_llm, load_env
from prompt_loader import render_prompt

WORD_SUBMITTER = "xxxxxx"


class OpeningContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    报告标题: str = Field(..., min_length=1)
    开篇亮明观点: str = Field(..., min_length=1)
    争议焦点段: str = Field(..., min_length=1)
    类案检索结论段: str = Field(..., min_length=1)


class SectionOneContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    第一部分标题: str = Field(..., min_length=1)
    第一部分正文: str = Field(..., min_length=1)


class SectionTwoContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    第二部分标题: str = Field(..., min_length=1)
    检索方法段: str = Field(..., min_length=1)
    类案检索情况段: str = Field(..., min_length=1)
    本案关联性段: str = Field(..., min_length=1)


class SectionThreeTableRow(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    中级人民法院案例: str = Field(..., min_length=1)
    案由: str = Field(..., min_length=1)
    裁判要点与理由: str = Field(..., min_length=1)


class SectionThreeViewpoint(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    观点标题: str = Field(..., min_length=1)
    观点总结: str = Field(..., min_length=1)
    表格标题: str = Field(..., min_length=1)
    裁判要点表格: list[SectionThreeTableRow] = Field(..., min_length=1)


class SectionThreeContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    第三部分标题: str = Field(..., min_length=1)
    第三部分引言段: str = Field(..., min_length=1)
    分类说明段: str = Field(..., min_length=1)
    观点列表: list[SectionThreeViewpoint] = Field(..., min_length=1)


class SectionFourLawItem(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    法规名称与条款: str = Field(..., min_length=1)
    条文原文: str = Field(..., min_length=1)
    条文要点列表: list[str] = Field(..., min_length=1)


class SectionFourContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    第四部分标题: str = Field(..., min_length=1)
    第四部分引言段: str = Field(..., min_length=1)
    法规条文列表: list[SectionFourLawItem] = Field(..., min_length=1)


class SectionFiveContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    第五部分标题: str = Field(..., min_length=1)
    分析结论列表: list[str] = Field(..., min_length=3, max_length=3)


class AttachmentCaseSelection(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    选中序号: list[int] = Field(..., min_length=1)


class WordReportContent(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    报告标题: str = Field(..., min_length=1)
    开篇亮明观点: str = Field(..., min_length=1)
    待决案件案情简述: str = Field(..., min_length=1)
    第一部分_类案基本事实概括: str = Field(..., min_length=1)
    第一部分_检索方法: str = Field(..., min_length=1)
    第一部分_检索情况: str = Field(..., min_length=1)
    第一部分_本案关联性: str = Field(..., min_length=1)
    第二部分_类案核心裁判要旨: str = Field(..., min_length=1)
    第二部分_观点总结列表: list[str] = Field(..., min_length=1)
    第四部分_相关法律法规原文: list[str] = Field(..., min_length=1)
    第五部分_结果分析: list[str] = Field(..., min_length=3, max_length=3)
    应当参照类案: list[str] = Field(..., min_length=1)
    可以参考类案: list[str] = Field(..., min_length=1)


async def generate_opening_content(query: str, rows: list[MinimalCaseRow]) -> OpeningContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    opening_writer = llm.with_structured_output(OpeningContent, method="function_calling")
    schema_text = json.dumps(OpeningContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows[:12]], ensure_ascii=False, indent=2)
    return await opening_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/opening_human.txt",
                    query=query,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_section_one_content(query: str, opening: OpeningContent, rows: list[MinimalCaseRow]) -> SectionOneContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    section_writer = llm.with_structured_output(SectionOneContent, method="function_calling")
    schema_text = json.dumps(SectionOneContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows[:12]], ensure_ascii=False, indent=2)
    return await section_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/section_one_human.txt",
                    query=query,
                    opening_viewpoint=opening.开篇亮明观点,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_section_two_content(query: str, section_one: SectionOneContent, rows: list[MinimalCaseRow]) -> SectionTwoContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    section_writer = llm.with_structured_output(SectionTwoContent, method="function_calling")
    schema_text = json.dumps(SectionTwoContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)
    return await section_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/section_two_human.txt",
                    query=query,
                    section_one_body=section_one.第一部分正文,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_section_three_content(query: str, section_two: SectionTwoContent, rows: list[MinimalCaseRow]) -> SectionThreeContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    section_writer = llm.with_structured_output(SectionThreeContent, method="function_calling")
    schema_text = json.dumps(SectionThreeContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)
    return await section_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/section_three_human.txt",
                    query=query,
                    section_two_summary=section_two.类案检索情况段,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_section_four_content(
    query: str,
    section_three: SectionThreeContent,
    rows: list[MinimalCaseRow],
) -> SectionFourContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    section_writer = llm.with_structured_output(SectionFourContent, method="function_calling")
    schema_text = json.dumps(SectionFourContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)
    return await section_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/section_four_human.txt",
                    query=query,
                    section_three_intro=section_three.第三部分引言段,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_section_five_content(
    query: str,
    section_three: SectionThreeContent,
    section_four: SectionFourContent,
    rows: list[MinimalCaseRow],
) -> SectionFiveContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    section_writer = llm.with_structured_output(SectionFiveContent, method="function_calling")
    schema_text = json.dumps(SectionFiveContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)
    return await section_writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/section_five_human.txt",
                    query=query,
                    section_three_intro=section_three.第三部分引言段,
                    section_four_intro=section_four.第四部分引言段,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def generate_word_report_content(query: str, rows: list[MinimalCaseRow]) -> WordReportContent:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")
    llm = init_llm()
    writer = llm.with_structured_output(WordReportContent, method="function_calling")
    schema_text = json.dumps(WordReportContent.model_json_schema(), ensure_ascii=False, indent=2)
    rows_text = json.dumps([row.model_dump(by_alias=True) for row in rows], ensure_ascii=False, indent=2)
    return await writer.ainvoke(
        [
            SystemMessage(content=render_prompt("word/writer_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/report_human.txt",
                    query=query,
                    rows_text=rows_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )


async def select_attachment_cases(query: str, rows: list[MinimalCaseRow], case_hits: list[CaseHit]) -> list[CaseHit]:
    if not query:
        raise ValueError("QUERY must not be empty.")
    if not case_hits:
        raise ValueError("case_hits must not be empty.")
    if len(rows) != len(case_hits):
        raise ValueError("rows and case_hits length mismatch.")
    llm = init_llm()
    selector = llm.with_structured_output(AttachmentCaseSelection, method="function_calling")
    schema_text = json.dumps(AttachmentCaseSelection.model_json_schema(), ensure_ascii=False, indent=2)

    candidates: list[dict[str, Any]] = []
    for idx, (row, hit) in enumerate(zip(rows, case_hits), start=1):
        candidates.append(
            {
                "序号": idx,
                "审理法院": hit.court_name,
                "案号": hit.case_no,
                "标题": hit.title,
                "案由": row.案由,
                "裁判结果": row.裁判结果,
                "主要原因": row.主要原因,
            }
        )
    candidates_text = json.dumps(candidates, ensure_ascii=False, indent=2)

    result: AttachmentCaseSelection = await selector.ainvoke(
        [
            SystemMessage(content=render_prompt("word/selector_system.txt")),
            HumanMessage(
                content=render_prompt(
                    "word/attachment_selector_human.txt",
                    query=query,
                    candidates_text=candidates_text,
                    schema_text=schema_text,
                )
            ),
        ]
    )

    selected_hits: list[CaseHit] = []
    seen: set[int] = set()
    for i in result.选中序号:
        if i < 1 or i > len(case_hits):
            raise ValueError(f"Attachment case index out of range: {i}")
        if i in seen:
            continue
        seen.add(i)
        selected_hits.append(case_hits[i - 1])
    if not selected_hits:
        raise ValueError("Selected attachment case count must be greater than 0.")
    return selected_hits


async def save_word(
    case_title: str,
    rows: list[MinimalCaseRow],
    attachment_case_hits: list[CaseHit],
    report_content: WordReportContent,
    report_title: str = "案涉争议问题检索报告",
    submitter: str = WORD_SUBMITTER,
    report_date: str | None = None,
) -> Path:
    if not case_title:
        raise ValueError("case_title must not be empty.")
    if not rows:
        raise ValueError("rows must not be empty.")

    def _build_and_save() -> Path:
        date_text = report_date or datetime.now().strftime("%Y年%m月")
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

        style = doc.styles["Normal"]
        style.font.name = "仿宋"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        style.font.size = Pt(14)

        label_color = RGBColor(47, 84, 150)

        def add_spacer(count: int) -> None:
            for _ in range(count):
                doc.add_paragraph("")

        def add_center_line(label: str, value: str, size: int) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"【{label}】")
            r1.font.name = "仿宋"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r1.font.bold = True
            r1.font.size = Pt(size)
            r1.font.color.rgb = label_color
            r2 = p.add_run(value)
            r2.font.name = "仿宋"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r2.font.bold = True
            r2.font.size = Pt(size)

        def add_body_paragraph(text: str) -> None:
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            for line in lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(28)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                run.font.size = Pt(12)

        def add_body_field_label(label: str) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(label)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(14)
            run.font.bold = True

        def add_body_field_value(label: str, value: str) -> None:
            value_lines = [line.strip() for line in value.splitlines() if line.strip()]
            for idx, line in enumerate(value_lines):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(28)
                p.paragraph_format.line_spacing = 1.5
                if idx == 0:
                    r1 = p.add_run(label)
                    r1.font.name = "仿宋"
                    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                    r1.font.size = Pt(12)
                    r1.font.bold = True
                r2 = p.add_run(line)
                r2.font.name = "仿宋"
                r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                r2.font.size = Pt(12)

        def add_section_heading(text: str) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = label_color

        def add_case_table(rows_data: list[CaseTableRow]) -> None:
            table = doc.add_table(rows=len(rows_data) + 1, cols=3)
            table.style = "Table Grid"

            headers = ["中级人民法院案例", "案由", "裁判要点与理由"]
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(header)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                run.font.size = Pt(12)
                run.font.bold = True

            for row_idx, row_data in enumerate(rows_data, start=1):
                values = [row_data.中级人民法院案例, row_data.案由, row_data.裁判要点与理由]
                for col_idx, value in enumerate(values):
                    cell = table.cell(row_idx, col_idx)
                    p = cell.paragraphs[0]
                    if col_idx < 2:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(value)
                    run.font.name = "仿宋"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                    run.font.size = Pt(12)

        def clean_attachment_text(text: str) -> str:
            cleaned_text = html.unescape(text)
            cleaned_text = cleaned_text.replace("\r\n", "\n").replace("\r", "\n")
            cleaned_text = re.sub(r"<[^>]+>", "", cleaned_text)
            cleaned_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned_text)
            cleaned_text = re.sub(r"__(.*?)__", r"\1", cleaned_text)
            cleaned_text = re.sub(r"`([^`]+)`", r"\1", cleaned_text)
            cleaned_text = re.sub(r"\\</?em>", "", cleaned_text)
            cleaned_text = re.sub(r"[ \t\f\v]+", " ", cleaned_text)
            cleaned_text = re.sub(r"\n\s*\n+", "\n", cleaned_text)
            return cleaned_text.strip()

        def add_attachment_title(text: str) -> None:
            title = clean_attachment_text(text)
            if not title:
                return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(title)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.underline = True

        def add_attachment_paragraph(text: str) -> None:
            cleaned_text = clean_attachment_text(text)
            lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
            for line in lines:
                p = doc.add_paragraph()
                is_case_no_line = bool(re.match(r"^（?\d{4}）", line))
                is_signature_line = bool(re.match(r"^(审判长|审判员|书记员|法官助理|代理审判员)", line))
                is_date_footer_line = bool(re.match(r"^(二〇|二○|二O|\d{4}年)", line))
                if is_case_no_line or is_signature_line or is_date_footer_line:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.first_line_indent = Pt(0)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line)
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                run.font.size = Pt(10.5)

        def add_attachment_field(label: str, value: str, right_aligned: bool = False) -> None:
            label_text = clean_attachment_text(label)
            value_text = clean_attachment_text(value)
            if not label_text and not value_text:
                return
            p = doc.add_paragraph()
            if right_aligned:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.first_line_indent = Pt(0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.line_spacing = 1.5
            r1 = p.add_run(label_text)
            r1.font.name = "仿宋"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r1.font.size = Pt(10.5)
            r1.font.bold = True
            r2 = p.add_run(value_text)
            r2.font.name = "仿宋"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            r2.font.size = Pt(10.5)

        add_spacer(1)
        add_center_line("案由", case_title, 16)
        add_spacer(4)
        add_center_line("标题", report_title, 22)
        add_spacer(7)
        add_center_line("署名", f"提交人：{submitter}", 14)
        add_spacer(1)
        add_center_line("日期", date_text, 14)

        doc.add_page_break()
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(report_content.报告标题)
        title_run.font.name = "仿宋"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        title_run.font.bold = True
        title_run.font.size = Pt(22)

        add_body_paragraph("尊敬的合议庭：")
        add_body_paragraph(report_content.开篇亮明观点)
        add_section_heading("【待决案件案情简述】")
        add_body_paragraph(report_content.待决案件案情简述)
        add_section_heading("一、【类案基本事实概括】")
        add_body_paragraph(report_content.第一部分_类案基本事实概括)
        add_body_field_value("【写明检索方法：包括检索平台、检索关键词、检索法院、检索时间】", report_content.第一部分_检索方法)
        add_body_field_value("【总结类案检索情况】", report_content.第一部分_检索情况)
        add_body_field_value("【要提及类案与本案直接的关联性】", report_content.第一部分_本案关联性)
        add_section_heading("二、【类案核心裁判要旨】")
        add_body_paragraph(report_content.第二部分_类案核心裁判要旨)
        for idx, view in enumerate(report_content.第二部分_观点总结列表, start=1):
            add_body_field_value(f"【观点{idx}】", view)
            viewpoint_table_rows = build_case_table_rows_for_viewpoint(rows, view, max_cases=3)
            add_case_table(viewpoint_table_rows)
        add_section_heading("三、【相关法律法规原文：法律、司法解释】")
        for idx, law_text in enumerate(report_content.第四部分_相关法律法规原文, start=1):
            add_body_paragraph(f"{idx}、{law_text}")
        add_section_heading("四、【类案检索结果分析】")
        for idx, item in enumerate(report_content.第五部分_结果分析, start=1):
            add_body_paragraph(f"{idx}、{item}")
        add_body_field_label("应当参照类案：")
        for idx, item in enumerate(report_content.应当参照类案, start=1):
            add_body_paragraph(f"{idx}、{item}")
        add_body_field_label("可以参考类案：")
        for idx, item in enumerate(report_content.可以参考类案, start=1):
            add_body_paragraph(f"{idx}、{item}")

        doc.add_page_break()
        add_section_heading("五、【附件：类案检索原文】")
        add_attachment_paragraph("案例原文：")
        for idx, hit in enumerate(attachment_case_hits, start=1):
            if idx > 1:
                doc.add_page_break()
            add_attachment_title(f"{hit.title}")
            add_attachment_field("审理法院：", f" {hit.court_name or ''}")
            add_attachment_field("（案号）", f"{hit.case_no or ''}", right_aligned=True)
            add_attachment_field("案由：", f" {hit.cause or ''}")
            add_attachment_field("裁判日期：", f" {hit.case_date or ''}")
            add_attachment_paragraph(hit.content)

        output_path = Path("outputs") / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    return await asyncio.to_thread(_build_and_save)


async def main() -> None:
    load_env()
    rows, case_hits = await build_rows()
    report_content = await generate_word_report_content(QUERY, rows)
    attachment_case_hits = await select_attachment_cases(QUERY, rows, case_hits)
    word_report_path = await save_word(
        case_title=QUERY,
        rows=rows,
        attachment_case_hits=attachment_case_hits,
        report_content=report_content,
    )
    print(f"Generated Word Report: {word_report_path.resolve()}")
    print(f"Rows: {len(rows)}")
    print(f"Attachment Cases: {len(attachment_case_hits)}")


if __name__ == "__main__":
    asyncio.run(main())
